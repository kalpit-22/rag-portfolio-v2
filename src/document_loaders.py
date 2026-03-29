import os
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import docx

def load_docx(file_path: str):
    """Custom simple loader for DOCX using python-docx."""
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return [Document(page_content=text, metadata={"source": file_path})]

def load_and_split_document(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200):
    """Loads a document (PDF, TXT, DOCX) and splits it into chunks."""
    ext = os.path.splitext(file_path)[-1].lower()
    
    if ext == '.pdf':
        loader = PyMuPDFLoader(file_path)
        docs = loader.load()
    elif ext == '.docx':
        docs = load_docx(file_path)
    elif ext == '.txt':
        loader = TextLoader(file_path, encoding='utf-8')
        docs = loader.load()
    else:
        raise ValueError(f"Unsupported file type for {file_path}")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    splits = text_splitter.split_documents(docs)
    return splits
